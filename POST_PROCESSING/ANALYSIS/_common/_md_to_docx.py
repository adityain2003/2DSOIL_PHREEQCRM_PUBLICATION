# -*- coding: utf-8 -*-
"""Minimal Markdown -> Word converter for this repo's docs.
Handles: ATX headings (#..####), paragraphs, '-' bullets, 'N.' numbered lists,
pipe tables (with header separator), fenced ``` code blocks, '>' blockquotes,
'---' rules, and inline `code` / **bold** / *italic*. House style: Times New Roman
12 pt, justified body; Consolas for code.  Usage: python _md_to_docx.py in.md out.docx
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = sys.argv[1]; OUT = sys.argv[2]
BODY = "Times New Roman"; MONO = "Consolas"
HSIZE = {1: 16, 2: 13, 3: 12, 4: 11}

doc = Document()
doc.styles["Normal"].font.name = BODY; doc.styles["Normal"].font.size = Pt(12)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.8); s.left_margin = s.right_margin = Inches(0.8)

def setf(r, font=BODY, size=12, bold=False, italic=False, color=None):
    r.font.name = font; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color

INLINE = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*\n]+\*)')
def inline(p, text, size=12, bold=False, italic=False, color=None):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            setf(p.add_run(text[pos:m.start()]), BODY, size, bold, italic, color)
        t = m.group(0)
        if t[0] == '`':
            setf(p.add_run(t[1:-1]), MONO, size-1, bold, italic, color)
        elif t.startswith('**'):
            setf(p.add_run(t[2:-2]), BODY, size, True, italic, color)
        else:
            setf(p.add_run(t[1:-1]), BODY, size, bold, True, color)
        pos = m.end()
    if pos < len(text):
        setf(p.add_run(text[pos:]), BODY, size, bold, italic, color)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3; p.paragraph_format.space_after = Pt(6)
    inline(p, text, 12, bold, italic); return p

def heading(text, lvl):
    wl = max(0, lvl - 1)                       # md '#'->Title(0), '##'->Heading 1, '###'->Heading 2
    style = "Title" if wl == 0 else "Heading %d" % min(wl, 9)
    p = doc.add_paragraph(style=style)         # real Word heading style -> navigation/outline/TOC
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    inline(p, text, HSIZE.get(lvl, 11), bold=True, color=RGBColor(0, 0, 0))

def listitem(text, ordered, n=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(3)
    setf(p.add_run((f"{n}. " if ordered else "• ")), BODY, 12, bold=ordered)
    inline(p, text, 12)

def codeblock(lines):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.rows[0].cells[0]
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_after = Pt(0)
        setf(p.add_run(ln if ln else " "), MONO, 9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def table(rows):
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            par = cells[ci].paragraphs[0]; par.paragraph_format.line_spacing = 1.0
            inline(par, txt, 10, bold=(ri == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def is_sep(line):
    s = line.strip().strip('|')
    return bool(s) and set(s.replace('|', '').replace(' ', '')) <= set(':-')

def cells(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]

lines = open(SRC, encoding="utf-8").read().split("\n")
i, n = 0, len(lines)
while i < n:
    ln = lines[i]
    st = ln.strip()
    if st.startswith("```"):                                  # fenced code
        i += 1; buf = []
        while i < n and not lines[i].strip().startswith("```"):
            buf.append(lines[i]); i += 1
        codeblock(buf); i += 1; continue
    if st.startswith("|") and i + 1 < n and is_sep(lines[i+1]):  # table
        rows = [cells(ln)]; i += 2
        while i < n and lines[i].strip().startswith("|"):
            rows.append(cells(lines[i])); i += 1
        table(rows); continue
    m = re.match(r'^(#{1,6})\s+(.*)$', st)
    if m:
        heading(m.group(2), len(m.group(1))); i += 1; continue
    if re.match(r'^(-{3,}|\*{3,})$', st):                      # hr
        doc.add_paragraph().paragraph_format.space_after = Pt(2); i += 1; continue
    if st.startswith(">"):                                     # blockquote
        para(st.lstrip(">").strip(), italic=True); i += 1; continue
    mnum = re.match(r'^(\d+)\.\s+(.*)$', st)
    if mnum:                                                   # numbered (join continuations)
        body = mnum.group(2); i += 1
        while i < n and lines[i].strip() and not re.match(r'^(\d+\.|[-*]|#|\||```|>)', lines[i].strip()):
            body += " " + lines[i].strip(); i += 1
        listitem(body, True, int(mnum.group(1))); continue
    if re.match(r'^[-*]\s+', st):                              # bullet (join continuations)
        body = re.sub(r'^[-*]\s+', '', st); i += 1
        while i < n and lines[i].strip() and not re.match(r'^(\d+\.|[-*]\s|#|\||```|>)', lines[i].strip()):
            body += " " + lines[i].strip(); i += 1
        listitem(body, False); continue
    if not st:                                                 # blank
        i += 1; continue
    buf = [st]; i += 1                                          # paragraph (join wrapped lines)
    while i < n and lines[i].strip() and not re.match(r'^(#|\||```|>|-{3,}|\d+\.\s|[-*]\s)', lines[i].strip()):
        buf.append(lines[i].strip()); i += 1
    para(" ".join(buf))

doc.save(OUT)
print("WROTE", OUT)
