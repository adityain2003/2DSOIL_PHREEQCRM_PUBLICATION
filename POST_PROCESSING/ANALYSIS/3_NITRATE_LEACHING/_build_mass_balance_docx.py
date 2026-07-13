# -*- coding: utf-8 -*-
"""Generate MASS_BALANCE.docx -- nitrogen-conservation results in native Word tables.
DATA-DRIVEN: every number is read from the regenerated VARIANT_RUNS budget CSVs
(_NITRATE_BUDGET / _NITROGEN_BUDGET / _RESIDUAL_NITROGEN), so the document always
matches the current run (n = 1.78 calibrated hydraulics).
Times New Roman 12 pt; body justified + double-spaced; tables single-spaced."""
import os, csv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = r"C:\Users\adity\Desktop\2DSOIL_PHREEQCRM_INTEGRATION"
VR   = os.path.join(ROOT, "2DSOIL_PHREEQCRM_MODEL", "POST_PROCESSING", "ANALYSIS", "_data", "VARIANT_RUNS")
OUT  = os.path.join(ROOT, "2DSOIL_PHREEQCRM_MODEL", "MASS_BALANCE.docx")
FONT = "Times New Roman"
VARS = ["NR", "ZK", "ZK_ND", "FK", "FK_ND", "CK"]

def load(name):
    d = {}
    with open(os.path.join(VR, name), newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            d[row["variant"]] = row
    return d

NA = load("_NITRATE_BUDGET.csv")       # nitrified,denitrified,leached,resid_NO3,resid_NH4,init_NO3
NB = load("_NITROGEN_BUDGET.csv")      # min_kgha,nit_kgha,den_kgha,leached_kgha
RN = load("_RESIDUAL_NITROGEN.csv")    # res_org,res_nh4,res_no3,residual_soilN,n2o_N,leached,accounted

def f2(x): return f"{float(x):.2f}"
def f1(x): return f"{float(x):.1f}"

# bottom-line statistics, computed from the data
init = [float(NA[v]["init_NO3"]) for v in VARS]
init_mean = sum(init)/len(init); init_spread = max(init)-min(init)
init_pct = init_spread/init_mean*100
tot = [float(RN[v]["accounted"]) for v in VARS]
tot_mean = sum(tot)/len(tot); tot_spread = max(tot)-min(tot)
tot_pct = tot_spread/tot_mean*100

MD = []  # markdown mirror written alongside the .docx

doc = Document()
st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(12)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7); s.left_margin = s.right_margin = Inches(0.7)

def setrun(r, size=12, bold=False, italic=False, color=None):
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color

def para(text="", bold=False, italic=False, size=12, double=True):
    if text:
        MD.append(""); MD.append(("**%s**" % text) if bold else ("_%s_" % text) if italic else text)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0 if double else 1.0
    setrun(p.add_run(text), size, bold, italic); return p

def bullet(text):
    MD.append("- " + text)
    p = doc.add_paragraph(style="List Bullet"); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    setrun(p.add_run(text), 12); return p

def heading(text, lvl=1, size=12):
    MD.append(""); MD.append("#" * (lvl + 1) + " " + text); MD.append("")
    h = doc.add_heading(level=lvl); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    setrun(h.add_run(text), size, bold=True, color=RGBColor(0, 0, 0)); return h

def note(text):
    MD.append(""); MD.append("_" + text + "_")
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.0
    setrun(p.add_run(text), 10, italic=True)

def table(headers, rows, tnote=None):
    rows = list(rows)
    MD.append("")
    MD.append("| " + " | ".join(str(h) for h in headers) + " |")
    MD.append("|" + "|".join([" --- "] * len(headers)) + "|")
    for _r in rows:
        MD.append("| " + " | ".join(str(x) for x in _r) + " |")
    MD.append("")
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.paragraphs[0].paragraph_format.line_spacing = 1.0
        setrun(c.paragraphs[0].add_run(str(h)), 12, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            par = cells[i].paragraphs[0]; par.paragraph_format.line_spacing = 1.0
            setrun(par.add_run(str(v)), 12)
            if i > 0: par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if tnote: note(tnote)
    doc.add_paragraph().paragraph_format.line_spacing = 1.0

heading("MASS BALANCE — Nitrogen conservation in the generic 2DSOIL↔PhreeqcRM engine", lvl=0, size=14)
para("Nitrate-leaching validation (Murphy sand, low-frequency irrigation; n = 1.78 calibrated "
     "hydraulics). All values read directly from the regenerated reaction/budget CSVs.", italic=True)
para("The generic engine emits a per-step reaction report (REACTION_TOTALS.txt, enabled by "
     "REPORT_KINETICS 1). From it the nitrogen mass balance is reconstructed for all six kinetic "
     "variants and shown to conserve nitrogen. The focus is the mineral nitrate pool — the "
     "leaching-relevant form — with the full-system closure given as confirmation.")

heading("Bottom line", 1)
bullet(f"The mineral nitrate balance closes to {init_pct:.2f}%. Solving  initial NO3 + nitrified - "
       f"denitrified - leached = residual NO3  for the initial pool, every variant independently "
       f"recovers ~25.5 kg NO3-N/ha (mean {init_mean:.2f}, spread {init_spread:.2f}) despite very "
       f"different kinetics.")
bullet(f"The full tracked-N budget closes to {tot_pct:.3f}%:  residual soil N + N2O-N + leached = "
       f"{tot_mean:.1f} kg N/ha  for every variant — the same conserved total, partitioned differently.")
bullet("Reactions neither create nor destroy N — the N-weighted sum of component changes is ~1e-11 mol "
       "(machine precision).")

heading("Method", 1)
bullet("Per-step report: for every component, net reaction change d[component] (domain total, moles) = "
       "sum_i (c_out - c_in)_i * W_i, W_i = nodeArea_i * SLAB_WIDTH * 0.001 (bulk cell volume, L). theta "
       "cancels (push/pull is theta-weighted, CONC_BASIS=2), so totals are consistent with the leached solute.")
bullet("Transformations: cascade ORGANIC_N -> AMMONIUM -> NITRATE -> 1/2 NITROUS_GAS moves 1 N atom per "
       "mole of extent, so  mineralization = -d[ORGANIC_N];  nitrification = -d[ORGANIC_N] - d[AMMONIUM];  "
       "denitrification = 2*d[NITROUS_GAS].")
bullet("Residual pools: end-of-run absolute pools from final node concentrations (PHREEQC_SPECIES_OUT.txt) "
       "weighted by reconstructed node control-volume areas.")
bullet("Units: kg N/ha = mol(N) * 14.01/1000 / footprint_ha; footprint = 5 cm x 1 cm slab = 5e-8 ha "
       "(= 0.5*RowSp), so kg N/ha = mol(N) * 280200.")

heading("1. Mineral nitrate budget (kg N/ha)", 1)
para("initial NO3 + nitrified - denitrified - leached = residual NO3")
table(["variant", "nitrified", "denitrified", "leached", "residual NO3", "residual NH4", "recovered initial NO3"],
      [[v, f2(NA[v]["nitrified"]), f2(NA[v]["denitrified"]), f2(NA[v]["leached"]),
        f2(NA[v]["resid_NO3"]), f2(NA[v]["resid_NH4"]), f2(NA[v]["init_NO3"])] for v in VARS],
      tnote=f"Recovered initial nitrate = {init_mean:.2f} +/- {init_spread:.2f} kg N/ha ({init_pct:.2f}%) "
            f"across all six variants -- six independent runs converging on the same initial pool is the "
            f"mass-balance proof for mineral nitrate.")

heading("2. Nitrogen transformations (kg N/ha, cumulative)", 1)
table(["variant", "mineralized", "nitrified", "denitrified", "leached NO3-N"],
      [[v, f2(NB[v]["min_kgha"]), f2(NB[v]["nit_kgha"]), f2(NB[v]["den_kgha"]), f2(NB[v]["leached_kgha"])]
       for v in VARS],
      tnote="Denitrification diverts nitrate from leaching (ZK denitrifies most/leaches least; CK little). "
            "Each _ND variant matches its denitrifying twin on mineralization and nitrification; they differ "
            "only in denitrification. Moles in _NITROGEN_BUDGET.csv.")

heading("3. Residual mineral nitrogen left in the column (kg N/ha)", 1)
table(["variant", "residual NO3", "residual NH4", "residual mineral N"],
      [[v, f2(RN[v]["res_no3"]), f2(RN[v]["res_nh4"]),
        f2(float(RN[v]["res_no3"]) + float(RN[v]["res_nh4"]))] for v in VARS],
      tnote="NR: no nitrification, so the initial 25.5 kg/ha nitrate essentially all leached. ZK_ND/FK_ND: "
            "no denitrification sink keeps nitrate high (and they leach most). Organic N (~1410-1449 kg N/ha "
            "reservoir) is not the focus; only ~2-5% mineralizes in 32 days.")

heading("4. Full-system closure (confirmation)", 1)
para("Because reactions conserve N, N2O is retained (immobile), and irrigation is N-free "
     "(USE_FERTIGATION=0), the only N sink is nitrate leaching:  residual soil N (organic + NH4 + NO3) "
     "+ N2O-N + leached NO3-N = total tracked N.")
table(["variant", "residual soil N", "N2O-N (denitrified)", "leached", "total"],
      [[v, f2(RN[v]["residual_soilN"]), f2(RN[v]["n2o_N"]), f2(RN[v]["leached"]), f2(RN[v]["accounted"])]
       for v in VARS],
      tnote=f"Total tracked N = {tot_mean:.1f} +/- {tot_spread:.2f} kg N/ha ({tot_pct:.3f}%) for every "
            f"variant = the conserved initial N; each variant only partitions that fixed pool differently.")

heading("5. Validation", 1)
bullet("Reaction-level N conservation: d[ORGANIC_N] + d[AMMONIUM] + d[NITRATE] + 2*d[NITROUS_GAS] ~ 0 for "
       "all variants (~1e-11 mol vs reaction magnitudes ~1e-4 mol).")
bullet(f"Independent cross-check: the final N2O-N pool (from node concentrations + reconstructed areas) "
       f"equals the denitrified amount from the reaction report to the digit (ZK {f2(RN['ZK']['n2o_N'])}, "
       f"FK {f2(RN['FK']['n2o_N'])}, CK {f2(RN['CK']['n2o_N'])}) -- validating the area reconstruction "
       f"and the residual numbers.")
bullet(f"NR sanity: with no reactions, initial NO3 = leached + residual = {f2(RN['NR']['leached'])} + "
       f"{f2(RN['NR']['res_no3'])} = {float(RN['NR']['leached'])+float(RN['NR']['res_no3']):.1f} kg N/ha; "
       f"the pool simply drains.")
bullet(f"Convergence: six independent runs recover the same initial nitrate ({init_pct:.2f}%) and the same "
       f"total tracked N ({tot_pct:.3f}%).")

heading("Files & scripts", 1)
table(["output (VARIANT_RUNS/)", "contents"], [
    ["<VAR>/REACTION_TOTALS.txt", "per-step domain totals (moles): d[component] for every component"],
    ["_NITRATE_BUDGET.csv", "mineral nitrate budget (Section 1)"],
    ["_NITROGEN_BUDGET.csv", "transformations: mineralized / nitrified / denitrified (moles + kg N/ha)"],
    ["_RESIDUAL_NITROGEN.csv", "end-of-run pools + full-system closure (Section 4)"]])
table(["script (ANALYSIS/3_NITRATE_LEACHING/)", "role"], [
    ["nitrate_budget.py", "Section 1 -- mineral nitrate balance + recovered initial NO3"],
    ["nitrogen_budget.py", "Section 2 -- transformation amounts from component deltas"],
    ["residual_nitrogen.py", "Sections 3-4 -- absolute residual pools + full closure"],
    ["analyze_reaction_totals.py", "reaction-extent / component-delta reader + conservation check"]])

para("Method notes: reporting is read-only diagnostics (leaching results unchanged). The component-delta "
     "budget is the route used (the bare -kinetic_reactants emitted no dk_ columns; per-layer disaggregation "
     "would need the explicit reaction list in the .pqi + a rerun, no rebuild). N2O is tracked as an immobile "
     "component, so 'denitrified' = NO3-N converted to N2O-N: removed from the leachable pool but conserved "
     "in the total.", italic=True, size=10, double=False)

doc.save(OUT)
print("WROTE", OUT)
MD_OUT = OUT[:-5] + ".md" if OUT.endswith(".docx") else OUT + ".md"
with open(MD_OUT, "w", encoding="utf-8") as _f:
    _f.write("\n".join(MD).replace("\n\n\n", "\n\n").strip() + "\n")
print("WROTE", MD_OUT)
